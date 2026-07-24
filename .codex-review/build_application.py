from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


OUT = Path('/Users/yangl/Documents/个人资料/赵一帆/TellerAssist/智晓通_参赛申报方案_重构版.docx')

# Resolved preset: narrative_proposal
# Named overrides:
# - cjk_font: Arial Unicode MS for broad Chinese glyph coverage across renderers.
# - proposal_brand: restrained CMB-inspired red for cover accents only.
# - architecture_tables: blue-gray layer fills for technical diagrams.
FONT_LATIN = 'Arial Unicode MS'
FONT_CJK = 'Arial Unicode MS'
BLUE = '2E74B5'
DARK_BLUE = '1F4D78'
INK = '0B2545'
LIGHT = 'F4F6F9'
BLUE_GRAY = 'E8EEF5'
GRAY = '666666'
LIGHT_GRAY = 'F2F4F7'
BORDER = 'CBD5E1'
CMB_RED = 'B51E2E'
CAUTION = '7A5A00'
CAUTION_FILL = 'FFF7DD'
RISK = '9B1C1C'
WHITE = 'FFFFFF'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in('w:tcBorders')
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
    for edge in ('top', 'start', 'bottom', 'end', 'insideH', 'insideV'):
        if edge not in kwargs:
            continue
        tag = f'w:{edge}'
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        for key, value in kwargs[edge].items():
            element.set(qn(f'w:{key}'), str(value))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=120):
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(total))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.find(qn('w:tblInd'))
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), str(indent))
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn('w:tcW'))
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(width))
            tc_w.set(qn('w:type'), 'dxa')
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size=None, bold=None, italic=None, color=None, east_asia=FONT_CJK):
    run.font.name = FONT_LATIN
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:ascii'), FONT_LATIN)
    rfonts.set(qn('w:hAnsi'), FONT_LATIN)
    rfonts.set(qn('w:eastAsia'), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('第 ')
    set_run_font(run, size=9, color=GRAY)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)
    run = paragraph.add_run(' 页')
    set_run_font(run, size=9, color=GRAY)


def add_numbering_definitions(doc):
    numbering = doc.part.numbering_part.element
    existing = [int(x.get(qn('w:abstractNumId'))) for x in numbering.findall(qn('w:abstractNum'))]
    abstract_id = max(existing or [0]) + 1

    def abstract(multilevel_type, fmt, text, left, hanging, font=None):
        nonlocal abstract_id
        abstract_el = OxmlElement('w:abstractNum')
        abstract_el.set(qn('w:abstractNumId'), str(abstract_id))
        multi = OxmlElement('w:multiLevelType')
        multi.set(qn('w:val'), multilevel_type)
        abstract_el.append(multi)
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        num_fmt = OxmlElement('w:numFmt')
        num_fmt.set(qn('w:val'), fmt)
        lvl.append(num_fmt)
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), text)
        lvl.append(lvl_text)
        jc = OxmlElement('w:lvlJc')
        jc.set(qn('w:val'), 'left')
        lvl.append(jc)
        ppr = OxmlElement('w:pPr')
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'num')
        tab.set(qn('w:pos'), str(left))
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), str(left))
        ind.set(qn('w:hanging'), str(hanging))
        ppr.append(ind)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '80')
        spacing.set(qn('w:line'), '290')
        spacing.set(qn('w:lineRule'), 'auto')
        ppr.append(spacing)
        lvl.append(ppr)
        if font:
            rpr = OxmlElement('w:rPr')
            rfonts = OxmlElement('w:rFonts')
            rfonts.set(qn('w:ascii'), font)
            rfonts.set(qn('w:hAnsi'), font)
            rpr.append(rfonts)
            lvl.append(rpr)
        abstract_el.append(lvl)
        numbering.append(abstract_el)
        used = abstract_id
        abstract_id += 1
        return used

    def instance(abs_id):
        existing_num = [int(x.get(qn('w:numId'))) for x in numbering.findall(qn('w:num'))]
        num_id = max(existing_num or [0]) + 1
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(num_id))
        abstract_ref = OxmlElement('w:abstractNumId')
        abstract_ref.set(qn('w:val'), str(abs_id))
        num.append(abstract_ref)
        numbering.append(num)
        return num_id

    bullet_abs = abstract('singleLevel', 'bullet', '•', 540, 280, 'Symbol')
    decimal_abs = abstract('singleLevel', 'decimal', '%1.', 540, 280)
    return instance(bullet_abs), instance(decimal_abs)


def set_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn('w:numPr'))
    if num_pr is None:
        num_pr = OxmlElement('w:numPr')
        ppr.append(num_pr)
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), '0')
    numid = OxmlElement('w:numId')
    numid.set(qn('w:val'), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(numid)


def keep_with_next(paragraph):
    ppr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement('w:keepNext')
    ppr.append(keep)


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement('w:cantSplit')
    tr_pr.append(node)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']
normal.font.name = FONT_LATIN
normal.font.size = Pt(11)
normal._element.rPr.rFonts.set(qn('w:ascii'), FONT_LATIN)
normal._element.rPr.rFonts.set(qn('w:hAnsi'), FONT_LATIN)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CJK)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

for name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 18, 10),
    ('Heading 2', 13, BLUE, 12, 6),
    ('Heading 3', 12, DARK_BLUE, 8, 4),
]:
    st = styles[name]
    st.font.name = FONT_LATIN
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st._element.rPr.rFonts.set(qn('w:ascii'), FONT_LATIN)
    st._element.rPr.rFonts.set(qn('w:hAnsi'), FONT_LATIN)
    st._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CJK)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True
    st.paragraph_format.line_spacing = 1.15

for style_name in ['Title', 'Subtitle']:
    st = styles[style_name]
    st.font.name = FONT_LATIN
    st._element.rPr.rFonts.set(qn('w:ascii'), FONT_LATIN)
    st._element.rPr.rFonts.set(qn('w:hAnsi'), FONT_LATIN)
    st._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CJK)

if 'Table Text' not in styles:
    table_style = styles.add_style('Table Text', WD_STYLE_TYPE.PARAGRAPH)
else:
    table_style = styles['Table Text']
table_style.font.name = FONT_LATIN
table_style.font.size = Pt(9.5)
table_style._element.rPr.rFonts.set(qn('w:ascii'), FONT_LATIN)
table_style._element.rPr.rFonts.set(qn('w:hAnsi'), FONT_LATIN)
table_style._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CJK)
table_style.paragraph_format.space_before = Pt(0)
table_style.paragraph_format.space_after = Pt(2)
table_style.paragraph_format.line_spacing = 1.15

bullet_num_id, decimal_num_id = add_numbering_definitions(doc)


def p(text='', *, bold=False, italic=False, color=None, size=None, align=None, before=0, after=8, line=1.333, keep=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line
    para.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(text)
    set_run_font(run, size=size or 11, bold=bold, italic=italic, color=color)
    if keep:
        keep_with_next(para)
    return para


def rich_p(parts, *, align=None, before=0, after=8, line=1.333, keep=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = line
    para.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    for item in parts:
        if isinstance(item, str):
            text, opts = item, {}
        else:
            text, opts = item
        set_run_font(para.add_run(text), size=opts.get('size', 11), bold=opts.get('bold'), italic=opts.get('italic'), color=opts.get('color'))
    if keep:
        keep_with_next(para)
    return para


def heading(text, level=1):
    return doc.add_paragraph(text, style=f'Heading {level}')


def bullet(text, *, bold_lead=None):
    para = doc.add_paragraph()
    set_num(para, bullet_num_id)
    if bold_lead and text.startswith(bold_lead):
        set_run_font(para.add_run(bold_lead), size=11, bold=True, color=INK)
        set_run_font(para.add_run(text[len(bold_lead):]), size=11)
    else:
        set_run_font(para.add_run(text), size=11)
    return para


def numbered(text):
    para = doc.add_paragraph()
    set_num(para, decimal_num_id)
    set_run_font(para.add_run(text), size=11)
    return para


def callout(label, text, fill=LIGHT, label_color=DARK_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, top={'val': 'single', 'sz': 6, 'color': BORDER}, bottom={'val': 'single', 'sz': 6, 'color': BORDER}, start={'val': 'single', 'sz': 16, 'color': label_color}, end={'val': 'single', 'sz': 6, 'color': BORDER})
    para = cell.paragraphs[0]
    para.style = table_style
    para.paragraph_format.space_after = Pt(0)
    set_run_font(para.add_run(label + ' '), size=10.5, bold=True, color=label_color)
    set_run_font(para.add_run(text), size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def table(headers, rows, widths, *, header_fill=LIGHT_GRAY, font_size=9.5, aligns=None):
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(t, widths)
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, header_fill)
        para = c.paragraphs[0]
        para.style = table_style
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(para.add_run(h), size=font_size, bold=True, color=INK)
    set_repeat_table_header(t.rows[0])
    for row_data in rows:
        row = t.add_row()
        cant_split(row)
        for i, val in enumerate(row_data):
            c = row.cells[i]
            para = c.paragraphs[0]
            para.style = table_style
            if aligns:
                para.alignment = aligns[i]
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(para.add_run(str(val)), size=font_size)
    for row in t.rows:
        for c in row.cells:
            set_cell_border(c, top={'val': 'single', 'sz': 4, 'color': BORDER}, bottom={'val': 'single', 'sz': 4, 'color': BORDER}, start={'val': 'single', 'sz': 4, 'color': BORDER}, end={'val': 'single', 'sz': 4, 'color': BORDER})
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


_page_break_count = 0


def page_break():
    """Keep only the cover break; let later sections flow to avoid orphan pages."""
    global _page_break_count
    _page_break_count += 1
    if _page_break_count == 1:
        doc.add_page_break()


# Running header/footer
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run_font(hp.add_run('智晓通｜AI+运营参赛申报方案'), size=9, color=GRAY)
footer = section.footer
fp = footer.paragraphs[0]
add_page_number(fp)


# Cover: proposal_centerpiece
p('招商银行杭州高新支行', bold=True, size=12, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, before=38, after=52, line=1.0)
p('“智晓通”', bold=True, size=30, color=CMB_RED, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, line=1.0)
p('对公业务可信多模态智能填单助手', bold=True, size=20, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=10, line=1.15)
p('——让AI可用、可控、可核验地进入柜面运营流程', size=12.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=50, line=1.15)

cover_table = doc.add_table(rows=4, cols=2)
set_table_geometry(cover_table, [1800, 7560])
cover_rows = [
    ('参赛赛道', 'AI+运营'),
    ('报送单位', '招商银行杭州高新支行'),
    ('参赛成员', '周苏卿 等（待核实）'),
    ('联系方式', '待补充'),
]
for idx, (label, value) in enumerate(cover_rows):
    left, right = cover_table.rows[idx].cells
    for c in (left, right):
        set_cell_border(c, bottom={'val': 'single', 'sz': 4, 'color': BORDER})
    set_cell_shading(left, LIGHT)
    lp = left.paragraphs[0]; lp.style = table_style
    rp = right.paragraphs[0]; rp.style = table_style
    set_run_font(lp.add_run(label), size=10, bold=True, color=DARK_BLUE)
    set_run_font(rp.add_run(value), size=10.5, color=INK)
p('第五届创造力大赛参赛方案｜2026年7月', size=10.5, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, before=58, after=2, line=1.0)

page_break()

heading('申报摘要', 1)
callout('项目定位', '“智晓通”不是替代柜员作出业务判断的自动办理系统，而是一套面向对公柜面的可信辅助工具：AI负责理解和建议，规则引擎负责校验，柜员保留最终确认与提交权。')
p('项目聚焦对公账户销户、网银管理员变更等资料多、字段长、重复录入多的场景，通过语音识别、证件OCR、结构化字段提取和业务规则校验，将客户证件、柜员语音及行内主数据汇聚为可追溯的“字段证据链”，实现从资料采集、智能预填、冲突提示到柜员确认的闭环。')

heading('核心价值主张', 2)
bullet('提效：减少企业名称、证件号码、地址等字段的重复抄录，缩短柜面资料录入及复核时间。', bold_lead='提效：')
bullet('控错：对统一社会信用代码、身份证号码、证件有效期和户名一致性执行确定性校验。', bold_lead='控错：')
bullet('可信：每个字段均保留来源、置信度、原始证据、校验结果和人工修改记录。', bold_lead='可信：')
bullet('可推广：通过“业务模板+字段字典+规则包+系统适配器”扩展到更多对公运营场景。', bold_lead='可推广：')

heading('申报信息', 2)
table(['项目要素', '申报内容'], [
    ('项目名称', '“智晓通”对公业务可信多模态智能填单助手'),
    ('首期场景', '对公账户销户申请辅助填单（生产试点阶段仅预填与校验）'),
    ('扩展场景', '网银管理员变更、单位证件更新、账户信息维护、印鉴变更等'),
    ('部署原则', '行内或经审批的专有环境部署；客户数据不进入公共模型服务'),
    ('责任边界', '系统提供辅助建议；柜员负责身份核实、业务审核与最终提交'),
], [2100, 7260])

callout('数据说明', '本文中的业务量、耗时、差错率和收益数据均为便于申报论证而设置的“测算值（待核实）”，不代表招商银行杭州高新支行真实经营数据。提交前应以支行抽样记录和系统统计进行替换。', fill=CAUTION_FILL, label_color=CAUTION)

page_break()

heading('一、项目背景与问题定义', 1)
heading('（一）场景背景', 2)
p('杭州高新技术产业集聚区内科创企业数量多、企业信息变更频繁，对公客户对办理效率和数字化体验具有较高期待。支行在服务企业客户时，对公账户销户、网银管理员变更、证件更新等业务仍存在资料反复核对、相同字段多次录入、长字符易错等共性问题。')
p('传统模式下，柜员需要在客户提交的营业执照、身份证件、申请书及行内系统之间反复切换，手工录入企业名称、统一社会信用代码、账号、法定代表人和经办人信息。业务高峰期，重复录入不仅延长客户等待时间，也占用柜员开展业务咨询和客户经营的时间。')

heading('（二）现状测算与核实口径', 2)
table(['指标', '建议测算值', '正式核实方法'], [
    ('相关业务量', '8～15笔/工作日（待核实）', '抽取连续4周柜面流水，按业务类型分组统计'),
    ('单笔资料录入与复核', '12～18分钟（待核实）', '至少采样30笔，记录开始、预填完成、提交三个时间点'),
    ('重复录入字段', '15～30个/笔（待核实）', '对照申请书、影像系统和柜面系统字段清单'),
    ('人工修改/返工率', '2%～5%（待核实）', '统计退回、补录、二次修改及内部差错登记'),
    ('高峰期客户等待', '20～40分钟（待核实）', '按时段抽取叫号及业务受理数据'),
], [1800, 2460, 5100], font_size=9.2)

heading('（三）需要解决的核心问题', 2)
bullet('效率问题：多个系统和纸质材料之间重复读取、抄录和复核。', bold_lead='效率问题：')
bullet('质量问题：证件号码、账号、企业名称等长字段人工录入易出现形近字和漏位。', bold_lead='质量问题：')
bullet('体验问题：客户无法感知内部录入过程，等待时间长且需多次确认。', bold_lead='体验问题：')
bullet('治理问题：现有录入过程对字段来源、修改原因和识别质量缺少统一量化记录。', bold_lead='治理问题：')

heading('（四）项目目标', 2)
callout('建设目标', '构建“感知采集—结构化理解—规则校验—人工确认—受控写入—全程审计”的可信闭环，在不改变柜员审核责任的前提下，提升对公柜面资料录入效率与质量。')

page_break()

heading('二、总体解决方案', 1)
heading('（一）设计原则', 2)
table(['原则', '落地要求'], [
    ('人机协同', 'AI只提供预填和提示；关键字段由柜员确认，禁止AI自主提交'),
    ('数据分源', '证件字段优先来自OCR，账户字段优先来自行内主数据，原因类字段可由语音采集'),
    ('规则兜底', '格式、一致性、必填项和业务准入由确定性规则执行，不交由大模型自由判断'),
    ('证据可溯', '字段保留来源、证据、置信度、校验结果、修改记录及确认人'),
    ('渐进接入', '先实现表单预填与模拟提交，再根据接口条件逐步接入业务系统'),
], [1700, 7660])

heading('（二）端到端业务流程', 2)
for step in [
    '柜员选择业务场景并建立一次性业务会话。',
    '高拍仪采集营业执照、身份证等材料，系统完成图像质量检测、证件分类和OCR。',
    '柜员通过按键说话方式补充业务类型、办理原因及其他非证件字段。',
    '系统将OCR、语音和行内主数据映射至统一字段模型。',
    '规则引擎执行格式校验、跨来源一致性校验、必填项校验和风险提示。',
    '柜员在证据对照界面查看低置信度、冲突和缺失字段并修正。',
    '柜员确认后生成待提交表单或标准业务报文；首期原型采用模拟提交。',
    '系统记录识别、修改、确认和提交结果，形成可审计、可评估的数据闭环。',
]:
    numbered(step)

heading('（三）柜员交互设计', 2)
p('工作台采用“左证据、右表单、上状态、下操作”的布局：左侧展示证件原图及识别区域，右侧展示业务表单；字段以绿色、黄色、红色分别表示已通过、待确认和冲突；柜员可点击任一字段回看来源。语音功能采用按键触发，避免持续监听及多人声音串扰。')

page_break()

heading('三、技术架构设计', 1)
heading('（一）总体架构', 2)
p('系统采用前后端分离、模型服务解耦、规则配置化和集成适配器隔离的架构。AI能力与核心业务系统之间设置可信控制层，任何模型输出均不能绕过校验与人工确认直接进入核心系统。')

arch = table(['接入与体验层', '智能能力层', '可信控制层', '集成与治理层'], [
    ('柜员工作台\n高拍仪/扫描仪\n定向麦克风', 'ASR语音识别\n证件分类与OCR\n结构化字段提取', '统一字段模型\n规则引擎\n置信度与冲突检测', '客户/账户查询\n表单与业务适配器\n审计、监控与密钥管理'),
], [2100, 2420, 2420, 2420], header_fill=BLUE_GRAY, font_size=9.2,
aligns=[WD_ALIGN_PARAGRAPH.CENTER] * 4)

heading('（二）核心组件', 2)
table(['组件', '主要职责', '关键控制'], [
    ('会话与流程编排', '维护一次业务会话、材料状态、字段状态和流程节点', '会话超时、幂等标识、异常恢复'),
    ('ASR与语义提取', '将柜员语音转换为业务意图及标准字段', '按键说话、结构化输出、低置信度不自动接受'),
    ('OCR与证件解析', '图像质检、证件分类、文字识别和版面字段抽取', '反光/模糊提示、证件规则校验、证据区域回显'),
    ('统一字段模型', '汇聚语音、OCR、主数据和人工录入结果', '保存字段来源、置信度、证据与修改记录'),
    ('规则引擎', '执行格式、跨源一致性、必填、业务和风险规则', '规则版本化、可配置、命中原因可解释'),
    ('表单与集成适配器', '将标准字段映射为申请表或业务接口报文', '字段白名单、人工确认、幂等和结果回查'),
    ('审计与运营分析', '记录全链路事件并生成效率、质量和模型指标', '日志脱敏、权限隔离、留存周期配置'),
], [1500, 4920, 2940], font_size=8.8)

heading('（三）推荐部署形态', 2)
bullet('行内部署：工作台、业务服务、规则引擎、模型网关和审计服务部署在行内或经审批的专有环境。', bold_lead='行内部署：')
bullet('模型隔离：通过统一模型网关调用经审批的ASR、OCR和语言模型服务，禁止业务应用直接访问公共模型。', bold_lead='模型隔离：')
bullet('数据最小化：原始音频默认实时处理后删除；证件影像复用现有影像管理策略，避免形成额外副本。', bold_lead='数据最小化：')
bullet('高可用与降级：模型服务不可用时保留人工录入通道，不影响柜面基本业务连续性。', bold_lead='高可用与降级：')

page_break()

heading('四、可信字段模型与智能处理方案', 1)
heading('（一）字段证据链', 2)
p('系统不把“识别结果”简单等同于“业务事实”。每个字段均形成独立证据对象，至少包含字段值、数据来源、原始证据定位、识别置信度、规则校验结果、人工修改前后值、确认人员和时间。')
table(['字段示例', '首选来源', '校验方式', '确认策略'], [
    ('企业名称', '营业执照OCR+行内客户主数据', '名称标准化后跨源比对', '不一致时强制人工确认'),
    ('统一社会信用代码', '营业执照OCR', '长度、字符集、校验位、主数据一致性', '校验全部通过方可自动预填'),
    ('对公账号', '行内账户查询', '账号存在性、户名、账户类型和状态', '不得仅由语音或OCR确认'),
    ('法定代表人证件号', '身份证OCR+客户档案', '身份证校验位、有效期、档案一致性', '掩码展示，关键字段人工确认'),
    ('办理原因', '柜员语音或人工选择', '同义词映射至标准字典', '低置信度改为下拉选择'),
], [1650, 2430, 3000, 2280], font_size=8.8)

heading('（二）置信度与人工兜底', 2)
table(['状态', '建议阈值', '系统行为'], [
    ('高置信度', '≥0.95且规则全部通过', '自动预填，仍在提交前统一确认'),
    ('中置信度', '0.80～0.95或存在轻微差异', '黄色提示，要求逐项核对'),
    ('低置信度', '<0.80或存在关键冲突', '不自动接受，要求重新采集或手工录入'),
    ('禁止自动确认', '账号、证件号、客户身份等关键字段', '无论置信度高低，均保留人工确认'),
], [1800, 2460, 5100])

heading('（三）语音处理策略', 2)
bullet('语音主要采集业务类型、办理原因、说明性内容和柜员操作指令。')
bullet('企业名称、证件号码、账号等长字符串优先来自OCR或行内主数据。')
bullet('语言模型必须按预定义JSON结构返回字段，不允许自由生成业务报文。')
bullet('同义词、简称和口语表达通过业务词典标准化，例如“一般户”映射为“一般存款账户”。')

heading('（四）OCR处理策略', 2)
bullet('先质检再识别：检测模糊、反光、遮挡、缺角和方向错误。')
bullet('先分类再抽取：根据营业执照、身份证等不同版式使用对应字段模板。')
bullet('先校验再预填：证件号、有效期及统一社会信用代码通过规则后进入表单。')
bullet('保留证据定位：柜员点击字段可回看证件中对应区域。')

page_break()

heading('五、系统接入与业务边界', 1)
heading('（一）核心系统接入策略', 2)
table(['优先级', '接入方式', '适用阶段', '评价'], [
    ('1', '标准业务API或服务总线', '生产建设', '接口稳定、审计清晰，为首选方式'),
    ('2', '柜面系统预填或插件机制', '试点/生产', '对现有操作习惯影响小'),
    ('3', '结构化报文或文件导入', '试点', '便于隔离，但需控制格式与重复导入'),
    ('4', '表单预览与模拟提交', '比赛原型', '无需改造核心，适合快速验证价值'),
    ('5', 'RPA模拟录入', '短期过渡', '页面变化敏感，不建议作为长期架构'),
], [850, 2750, 1650, 4110], font_size=9.0,
aligns=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])

heading('（二）业务责任边界', 2)
callout('硬性边界', '“智晓通”不替代客户身份核实、业务准入判断、授权审批和柜员审核，不直接操作核心数据库，不在未经人工确认的情况下提交任何业务。', fill='FDECEC', label_color=RISK)
bullet('AI模块输出只能进入统一字段模型，不能直连核心系统。')
bullet('规则引擎负责确定性校验，但不替代现行业务制度及授权流程。')
bullet('柜员在提交前必须完成关键字段确认，系统记录确认人和确认时间。')
bullet('电子签名、电子凭证和无纸化归档属于二期范围，需独立完成制度及技术评审。')

heading('（三）异常与降级流程', 2)
table(['异常场景', '系统处置', '柜员动作'], [
    ('证件反光或模糊', '停止字段接受并提示重新采集', '调整角度或改用人工录入'),
    ('OCR与主数据冲突', '标红冲突字段并展示双方证据', '核实原件和客户档案'),
    ('语音识别不确定', '保留转写并弹出标准选项', '选择或手工修正'),
    ('模型服务超时', '结束智能处理但保留已采集材料', '切回传统人工流程'),
    ('业务接口失败', '不重复提交，记录幂等号并查询结果', '根据回查结果继续或转人工'),
], [1920, 4080, 3360], font_size=9.0)

page_break()

heading('六、安全、合规与审计设计', 1)
heading('（一）数据安全', 2)
bullet('部署环境：客户信息仅在行内网络或经审批的专有环境内处理。')
bullet('传输与存储：全链路加密，敏感字段按分类分级要求加密或脱敏。')
bullet('数据最小化：仅采集当前业务所需字段，原始音频默认不长期保留。')
bullet('训练隔离：生产客户数据不得未经审批用于模型训练；测试优先使用脱敏或合成数据。')
bullet('日志保护：应用日志不得记录完整账号、证件号、客户姓名或原始模型提示内容。')

heading('（二）访问控制与审计', 2)
table(['控制域', '控制措施', '留痕内容'], [
    ('身份与权限', '接入柜员统一认证，按岗位、机构、业务场景最小授权', '登录、查询、采集、修改、确认、提交'),
    ('字段操作', '关键字段掩码展示，查看完整信息需满足业务权限', '原值摘要、修改后值、修改原因、操作人'),
    ('模型调用', '经模型网关统一路由、限流和策略控制', '模型版本、请求类型、耗时、置信度、错误码'),
    ('规则执行', '规则版本化发布，可追溯命中原因', '规则编号、版本、输入摘要、执行结果'),
    ('业务提交', '双重校验、幂等控制、结果回查', '业务流水号、提交人、时间、返回状态'),
], [1740, 4320, 3300], font_size=8.8)

heading('（三）模型风险控制', 2)
bullet('输出约束：限定业务意图和字段枚举，采用结构化模式校验。')
bullet('事实校验：模型不得自行推断缺失的证件号、账号、日期或客户身份信息。')
bullet('版本治理：模型和提示模板变更须经过离线回归测试和灰度验证。')
bullet('持续监控：按业务、证件类型、字段类别统计准确率、人工修改率和漂移情况。')
bullet('人工可撤销：柜员可拒绝全部AI结果并立即切换为传统流程。')

page_break()

heading('七、MVP实施方案与推广路径', 1)
heading('（一）比赛原型范围', 2)
callout('首期场景', '选取“对公账户销户申请辅助填单”作为比赛演示场景，完成资料采集、语音补充、智能预填、规则校验、证据回看、柜员确认、模拟提交和指标看板。生产试点阶段仅承担资料预填与校验，不替代销户资格判断或核心记账。')

table(['纳入范围', '暂不纳入范围'], [
    ('营业执照、身份证OCR及图像质检', '销户资格的自动审批'),
    ('业务类型、账户类型、办理原因的语音提取', 'AI自主提交或直接写入核心数据库'),
    ('统一字段模型、规则校验和证据回看', '电子签名及全流程无纸化'),
    ('柜员确认、表单生成和模拟提交', '跨机构大规模推广'),
    ('效率、准确率、人工修改率看板', '使用真实客户数据进行开放式模型训练'),
], [4680, 4680], header_fill=BLUE_GRAY)

heading('（二）实施阶段', 2)
table(['阶段', '周期建议', '主要工作', '关键产出'], [
    ('0.业务梳理', '1～2周', '字段清单、流程、样本、规则和数据口径', '业务蓝图、字段字典、样本集'),
    ('1.原型开发', '4～6周', 'OCR、语音、预填、校验、确认和模拟提交', '可演示MVP、离线测试报告'),
    ('2.联调试点', '6～8周', '身份、主数据、影像或业务接口接入', '试点版本、安全与审计方案'),
    ('3.灰度运营', '4～8周', '1～2个网点试运行、A/B对比、问题整改', '试点评估报告、推广决策'),
], [1350, 1350, 3900, 2760], font_size=8.8)

heading('（三）项目团队建议', 2)
table(['角色', '建议投入', '职责'], [
    ('业务负责人/产品经理', '1人', '业务边界、流程、字段规则和试点协调'),
    ('前端工程师', '1人', '柜员工作台、设备交互、证据对照界面'),
    ('后端工程师', '1～2人', '流程编排、规则、数据模型、接口和审计'),
    ('AI/OCR工程师', '1人', 'ASR、OCR、结构化提取及评测'),
    ('测试工程师', '1人', '功能、回归、性能、异常与业务验收'),
    ('安全/运维/核心系统人员', '按阶段参与', '合规评审、环境、监控和系统联调'),
], [2340, 1620, 5400])

page_break()

heading('八、试点指标与效益测算', 1)
heading('（一）验收指标', 2)
table(['指标', '建议目标值', '统计口径'], [
    ('单笔资料录入与复核时间', '较基线降低≥30%', '同类业务、同等材料完整度的前后对比'),
    ('高置信度字段自动填充覆盖率', '≥80%', '满足阈值且规则通过的预填字段/全部可预填字段'),
    ('自动接受字段准确率', '≥99%', '无需人工修改的自动接受字段/自动接受字段'),
    ('关键字段人工确认率', '100%', '已人工确认关键字段/全部关键字段'),
    ('字段规则校验执行率', '100%', '已执行规则的适用字段/全部适用字段'),
    ('表单退回或补录率', '较基线降低≥30%', '因录入或材料问题发生的退回、补录业务占比'),
    ('OCR单份证件处理时间', '≤3秒', '图像上传完成至结构化字段返回'),
    ('语音结束至字段呈现时间', '≤2秒', '结束说话至结构化字段展示'),
    ('AI自主提交率', '0%', '未经柜员确认直接提交的业务数量'),
    ('全流程可追溯率', '100%', '具备完整操作与字段证据日志的业务占比'),
], [2600, 2160, 4600], font_size=8.5)

heading('（二）支行级年度效益测算示例', 2)
p('以下仅用于展示测算方法，正式申报时应替换为支行核实数据。')
table(['测算项', '示例值（待核实）', '计算说明'], [
    ('相关业务量', '10笔/工作日', '销户、网银管理员变更、证件更新等合计'),
    ('年工作日', '250日', '按实际营业安排调整'),
    ('单笔节省时间', '6分钟', '基线15分钟，试点后9分钟'),
    ('年度节省柜员工时', '约250小时', '10×250×6÷60'),
    ('相当于释放工作量', '约31个工作日', '按8小时/工作日折算'),
    ('质量收益', '暂不折算金额', '以补录率、退回率和差错整改工时衡量'),
], [2340, 2340, 4680])

heading('（三）评估设计', 2)
bullet('基线期：连续采集至少2周传统流程数据。')
bullet('试点期：在相同业务类型和材料完整度条件下采集至少4周数据。')
bullet('样本量：每个首期场景建议不少于100笔有效业务或脱敏回放样本。')
bullet('分层分析：分别统计证件类型、字段类型、柜员熟练度和高峰/非高峰时段。')
bullet('主观评价：收集柜员易用性、客户等待感知和异常处置便利度。')

page_break()

heading('九、创新性、独特价值与赛事契合', 1)
heading('（一）项目创新点', 2)
table(['创新点', '区别于普通OCR填单的价值'], [
    ('多模态字段分源', '根据字段风险选择OCR、语音、主数据或人工作为首选来源，而非“一种模型填全部”'),
    ('字段级证据链', '每个字段均可回看来源、原始证据、置信度和修改记录'),
    ('AI与规则双引擎', 'AI负责理解非结构化内容，规则负责确定性校验和风险拦截'),
    ('人在回路', '关键字段强制人工确认，AI不自主办理、不自主提交'),
    ('模板化推广', '通过字段字典、规则包和适配器扩展业务，而非为每个场景重建系统'),
    ('运营数据闭环', '识别质量、人工修改和流程耗时可量化，为运营优化提供依据'),
], [2300, 7060])

heading('（二）与“AI+运营”赛道的结合', 2)
bullet('AI+交互：通过ASR和结构化语义提取降低柜员操作负担。')
bullet('AI+识别：通过证件OCR和图像质量检测减少人工抄录。')
bullet('AI+流程：将资料采集、预填、核验和提交组织为可控闭环。')
bullet('AI+风控：以规则、置信度、证据回看和人工确认约束模型风险。')
bullet('AI+运营分析：基于流程和字段日志识别效率瓶颈及差错热点。')

heading('（三）推广价值', 2)
p('项目一旦形成统一字段模型和可信控制底座，新场景主要通过配置业务模板、字段映射和规则包完成扩展，可由单一支行试点逐步推广至更多网点。高新支行兼具科创企业客户集中、客户数字化体验要求高和业务变化较快等特点，具备较好的先行验证价值。')

callout('项目差异化表述', '“智晓通”构建的不是单一语音或OCR工具，而是一套具备字段证据、置信度、业务规则校验和人工确认机制的可信多模态填单体系。')

page_break()

heading('十、主要风险与应对措施', 1)
table(['风险', '影响', '应对措施', '责任建议'], [
    ('核心系统接口暂不可用', '无法形成真实提交闭环', '原型采用模拟提交；同步申请标准接口或预填机制', '业务与科技联合'),
    ('柜面噪声影响ASR', '语音转写错误或体验下降', '按键说话、定向麦克风、业务词典、下拉兜底', 'AI与前端'),
    ('证件图像质量不稳定', 'OCR错误或字段缺失', '图像质检前置、重拍提示、证据对照', 'AI与业务'),
    ('数据与隐私风险', '客户信息泄露或不当留存', '行内部署、加密脱敏、最小化采集、明确留存周期', '安全与运维'),
    ('模型版本漂移', '准确率下降或输出结构变化', '固定版本、回归样本、灰度发布、指标监控', 'AI与测试'),
    ('柜员过度依赖AI', '核验责任弱化', '关键字段强制确认、培训、抽检和责任提示', '业务管理'),
    ('收益难以证明', '比赛或推广说服力不足', '提前定义基线、采样方法和A/B评估口径', '产品与运营'),
], [1500, 2100, 4200, 1560], font_size=8.3)

heading('十一、结论与申报建议', 1)
p('“智晓通”立足对公柜面真实痛点，将语音识别、证件OCR和结构化理解能力嵌入现有运营流程，并通过规则校验、字段证据链和柜员确认机制解决AI在银行生产场景中的可信使用问题。项目首期范围清晰、演示闭环完整、技术路径可行，具备从支行原型向网点试点和多业务复制推广的潜力。')
p('建议以“表单预填与可信校验”为首期建设边界，以“对公账户销户申请辅助填单”为比赛展示场景，在完成真实数据核实、接口可行性评估和安全评审后，逐步扩展至网银管理员变更、证件更新和账户信息维护等场景。')

page_break()

heading('附录A：提交前待核实材料清单', 1)
table(['序号', '待核实材料', '建议责任人', '用途'], [
    ('1', '近4周相关业务量及业务类型分布', '运营主管', '确定场景规模和年度收益'),
    ('2', '至少30笔业务的录入、复核和总办理时长', '柜面业务人员', '建立效率基线'),
    ('3', '退回、补录、二次修改及内部差错记录', '运营主管', '建立质量基线'),
    ('4', '两类首期业务的完整表单与字段清单', '业务专家', '形成字段字典和表单模板'),
    ('5', '脱敏证件与语音样本各100份以上', '业务/科技', '开展离线识别评测'),
    ('6', '客户、账户、影像及柜面系统接口清单', '科技联系人', '确认接入路径'),
    ('7', '数据分类分级、留存及日志脱敏要求', '安全合规人员', '完成安全设计'),
    ('8', '参赛成员、联系方式、赛制字数和附件要求', '项目负责人', '完成最终报送'),
], [700, 4000, 1800, 2860], font_size=8.8)

heading('附录B：比赛演示脚本建议', 1)
for text in [
    '展示传统痛点：柜员需要在营业执照、身份证、申请书和系统间反复切换。',
    '扫描营业执照和身份证，系统显示图像质检及OCR字段。',
    '柜员说：“办理一般户销户，原因是企业迁址。”系统提取标准业务字段。',
    '系统从行内模拟主数据读取账户信息，并对企业名称和证件信息进行交叉校验。',
    '故意设置一处证件信息冲突，展示红色提示、证据回看和人工修正。',
    '柜员确认表单并执行模拟提交，系统生成电子申请表及审计记录。',
    '看板展示单笔耗时、自动填充率、人工修改率和规则拦截情况。',
]:
    numbered(text)

heading('附录C：建议用于最终申报的关键表述', 1)
callout('一句话介绍', '“智晓通”是一套面向对公柜面的可信多模态智能填单助手，通过语音、OCR和行内主数据协同采集信息，以规则校验、字段证据链和柜员确认保障AI结果可用、可控、可核验。')
callout('核心边界', 'AI负责理解与建议，规则引擎负责校验，柜员保留最终确认与提交权。')
callout('推广模式', '沉淀统一字段模型、业务模板、规则包和系统适配器，由单场景验证逐步扩展至多业务、多网点。')

# Core properties
doc.core_properties.title = '“智晓通”对公业务可信多模态智能填单助手参赛申报方案'
doc.core_properties.subject = '第五届创造力大赛 AI+运营赛道'
doc.core_properties.author = '招商银行杭州高新支行项目组'
doc.core_properties.keywords = '智晓通, AI+运营, 对公业务, OCR, 语音识别, 智能填单'

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
